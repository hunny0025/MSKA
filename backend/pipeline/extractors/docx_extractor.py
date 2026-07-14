"""
DOCX text extractor.
"""
from pipeline.extractors import ExtractionError


def extract(file_path: str) -> str:
    """Extract text from a .docx file including tables. Raises ExtractionError on failure."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ExtractionError("python-docx is not installed")

    try:
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        result = "\n".join(parts).strip()
        if not result:
            raise ExtractionError(f"DOCX produced empty text: {file_path}")
        return result
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"DOCX extraction failed for {file_path}: {exc}") from exc
